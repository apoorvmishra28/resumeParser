# import textract
#
# def extract_skills(file):
#     if 'skills' in file:
#         import pdb; pdb.set_trace()
#         skills = file
#
# text = textract.process('/home/cis/Documents/CV/CV_MACHINE_LEARNING.doc')
#
# skills = extract_skills(text)
#
# print(skills)
#
#


# import tensorflow as tf

import docx
# import nltk
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
import string

# document = docx.Document('CV_MACHINE_LEARNING.docx')
# document = docx.Document('Resume1.docx')
document = docx.Document('executive_cv.docx')

# nltk.download()

ps = PorterStemmer()

tables = document.tables
# import pdb; pdb.set_trace()
experience = []
skills = []


for table in tables:
    for row in table.rows:
        for cell in row.cells:
            words = word_tokenize(cell.text)
            for w in words:
                # import pdb; pdb.set_trace()
                if 'experi' in ps.stem(w):
                    for paragraph in cell.paragraphs:
                        experience.append(paragraph.text)
                        print(paragraph.text)
                if 'skill' in ps.stem(w):
                    for paragraph in cell.paragraphs:
                        skills.append(paragraph.text)
                        print(paragraph.text)
            break


for paragraph in document.paragraphs:
    words = word_tokenize(paragraph.text)
    import pdb; pdb.set_trace()
    for word in words:
        if 'experi' in ps.stem(word):
            experience.append(paragraph.text)
            print paragraph.text

print experience
print skills
